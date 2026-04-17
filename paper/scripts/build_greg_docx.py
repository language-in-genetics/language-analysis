#!/usr/bin/env python3

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "paper" / "draft-sections" / "greg-baker-methods-results.docx"
FIG1 = ROOT / "paper" / "figures" / "figure-1-focal-term-trends.png"
FIG2 = ROOT / "paper" / "figures" / "figure-2-corpus-scope.png"


def add_paragraph(document: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    para = document.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Times New Roman"


def main() -> None:
    document = Document()
    style_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Greg Baker Draft: Methods, Results, and Figure Text")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Prepared as a separate merge document to match the public dashboard at lig.symmachus.org "
        "as checked on 2026-04-17."
    )
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    document.add_heading("Materials and Methods", level=1)

    document.add_heading("Title and Abstract Analysis", level=2)
    add_paragraph(
        document,
        "We assembled article metadata from CrossRef for the 17 human genetics journals enabled in the "
        "project database at the time of analysis. The in-scope corpus contained 98,699 records published "
        "between 1947 and 2025. Of these, 96,489 records (97.8%) had been processed by the automated "
        "classifier when the present export was generated, while 2,243 records were skipped because no "
        "title was available. Processed records included 26,221 articles with abstracts and 70,268 "
        "title-only records."
    )
    add_paragraph(
        document,
        "Each record was classified from its title and, when available, its abstract using an OpenAI "
        "batch-processing pipeline with structured function-calling output. The classifier was prompted "
        "to determine whether the article used terms such as \"Caucasian\", \"white\", or "
        "\"European ancestry\" in a way that referred to race, ancestry, ethnicity, or population. "
        "The stored outputs were article-level boolean fields for caucasian, white, and european, "
        "together with the extracted phrase where relevant. Articles were allowed to be positive for "
        "more than one category, but each category was counted at most once per article regardless of "
        "repetition within the title or abstract."
    )
    add_paragraph(
        document,
        "Although the original research question referred to the phrase \"European ancestry\", the "
        "implemented european field captured a broader class of European-origin descriptors rather than "
        "the exact phrase alone. In the processed corpus, this category included phrases such as "
        "\"European populations\", \"European ancestry\", \"European Americans\", \"Europeans\", and "
        "\"European descent\"."
    )

    document.add_heading("Data Validation", level=2)
    add_paragraph(
        document,
        "Because white is the most ambiguous of the focal categories, we conducted a manual spot-check "
        "of 80 randomly sampled white-positive records from the processed corpus. Seventy-eight of the "
        "80 sampled records were clear references to human race, ethnicity, ancestry, or population. "
        "Two were clear false positives: one Heredity article on the Drosophila white gene and one case "
        "report using white to describe hypopigmented skin. This spot-check suggests that obvious false "
        "positives are uncommon but not absent, and that the remaining errors are concentrated in "
        "non-human genetics and colour-description contexts."
    )

    document.add_heading("Data Analysis", level=2)
    add_paragraph(
        document,
        "The primary outcome was article-level prevalence rather than raw token frequency. For each "
        "publication year, we calculated the proportion of processed articles that were positive for each "
        "focal category. To reduce year-to-year noise, especially in early decades with smaller yearly "
        "volumes, we overlaid the annual percentages with centred 5-year moving averages. Because 2025 "
        "processing was incomplete at export time, all temporal figures and trend statements were "
        "restricted to 1947-2024."
    )

    document.add_heading("Results", level=1)

    document.add_heading("Corpus Coverage", level=2)
    add_paragraph(
        document,
        "At export time, the project database contained 98,699 in-scope records from 17 journals, of "
        "which 96,489 had been processed by the title/abstract classifier and 2,243 had been skipped "
        "because no title was available. After excluding the incomplete 2025 tail, the trend analysis "
        "covered 95,941 processed records published between 1947 and 2024. Only 26,027 of these 95,941 "
        "records (27.3%) included an abstract in CrossRef, meaning that the majority of classifications "
        "relied on titles alone."
    )
    add_paragraph(
        document,
        "Across 1947-2024, the three focal categories appeared in 1,667 processed records (1.7%). "
        "A broader set of population descriptors, including the exploratory other category, appeared in "
        "5,188 records (5.4%)."
    )

    document.add_heading("Temporal Trends in Focal Terminology", level=2)
    add_paragraph(
        document,
        "Across the 1947-2024 series, European-origin descriptors were the most common focal category "
        "(933 articles), followed by Caucasian (530) and white (282). The term Caucasian increased "
        "through the 1990s and early 2000s, peaking at 1.42% of processed articles in 2007 (36/2,539), "
        "then declining to 0.13% in 2024 (3/2,328). Its centred 5-year moving average peaked in 2006 at "
        "1.22% and fell steadily thereafter."
    )
    add_paragraph(
        document,
        "White remained relatively uncommon through the earlier decades, then rose in the late 2010s and "
        "early 2020s. The annual percentage peaked at 0.74% in 2020 (20/2,685) and remained 0.73% in "
        "2021 (20/2,736). The centred 5-year moving average peaked in 2022 at 0.65% and remained "
        "elevated at 0.59% in 2024."
    )
    add_paragraph(
        document,
        "European-origin descriptors were both more frequent and more persistent than the other two focal "
        "categories. The annual percentage peaked at 1.83% in 2021 (50/2,736). The centred 5-year moving "
        "average continued to rise into the most recent complete years, peaking at 1.56% in 2023 and "
        "remaining high at 1.46% in 2024. The most common stored formulations in this category were "
        "\"European populations\" (n = 52), \"European ancestry\" (n = 33), \"European Americans\" "
        "(n = 21), \"Europeans\" (n = 21), and \"European descent\" (n = 15), indicating that the "
        "category captured a broader European-origin vocabulary rather than the exact phrase "
        "\"European ancestry\" alone."
    )
    add_paragraph(
        document,
        "Taken together, these patterns suggest a shift away from Caucasian, a later but smaller rise in "
        "white, and sustained growth in European-origin language."
    )

    document.add_heading("Manual Spot-Check of white", level=2)
    add_paragraph(
        document,
        "In the manual review of 80 randomly sampled white-positive records, 78 were judged to be true "
        "positives and 2 were clear false positives. The false positives were a Heredity paper about the "
        "Drosophila white gene and a clinical case report that used white as a colour descriptor for skin "
        "changes. These findings support the use of the title/abstract classifier for descriptive trend "
        "analysis, while also showing that ambiguous uses of white are not fully eliminated."
    )

    document.add_heading("Figure Legends", level=1)
    add_paragraph(
        document,
        "Figure 1. Article-level prevalence of focal terminology in titles and abstracts, 1947-2024. "
        "Each panel shows the annual percentage of processed articles classified as using Caucasian, "
        "white, or a broader European-origin descriptor. Thin lines show annual values; thick lines show "
        "centred 5-year moving averages. The 2025 data were excluded because processing was incomplete "
        "at export time."
    )
    add_paragraph(
        document,
        "Figure 2. Corpus scope by journal at export time. Bar length shows the number of in-scope "
        "records per journal, and bar colour shows the percentage of records for which CrossRef supplied "
        "an abstract. This figure illustrates the uneven abstract coverage across journals and explains "
        "why the classifier frequently operated on titles alone."
    )

    document.add_page_break()
    document.add_heading("Figures", level=1)

    figure_1_label = document.add_paragraph("Figure 1")
    figure_1_label.paragraph_format.keep_with_next = True
    document.add_picture(str(FIG1), width=Inches(6.5))

    document.add_page_break()
    figure_2_label = document.add_paragraph("Figure 2")
    figure_2_label.paragraph_format.keep_with_next = True
    document.add_picture(str(FIG2), width=Inches(6.5))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT)


if __name__ == "__main__":
    main()
